#!/usr/bin/env python3
"""
Strip bracketed guidance placeholders from value cells in BAD Marketing brief templates.

The templates ship with helper text like "[the master picture folder for this brief]"
in the second column of overview/creative tables. This text is useful when a human
opens the template directly, but should not appear in a generated brief.

This tool walks tables 0..5, looks at the second cell of each non-header row,
and clears any cell whose entire content matches the bracket-placeholder pattern.

Usage:
  python3 strip_placeholders.py templates/static.docx
  python3 strip_placeholders.py templates/static.docx templates/static.clean.docx
"""

import sys
import os
import re
import shutil
from docx import Document
from docx.oxml.ns import qn

# Match cells whose entire content (after stripping whitespace) is one or more
# bracketed phrases. Examples to match:
#   "[Any general notes]"
#   "[Specific info about the avatar]Example: mid-high income corporate men/business..."
#   "[Any specific design notes]"
# We're intentionally permissive: if the cell text STARTS with "[", treat it as guidance.
BRACKET_RE = re.compile(r'^\s*\[.+', re.DOTALL)


def cell_text(cell):
    return ''.join(p.text for p in cell.paragraphs)


def clear_cell(cell):
    """Remove all paragraphs except the first, and clear the first paragraph's runs."""
    # Remove extra paragraphs
    paras = cell.paragraphs
    for p in paras[1:]:
        p._element.getparent().remove(p._element)
    # Clear runs in the first paragraph (preserve pPr)
    first = paras[0]
    for child in list(first._element):
        if child.tag != qn('w:pPr'):
            first._element.remove(child)


def strip_placeholders(input_path, output_path):
    if input_path != output_path:
        shutil.copy2(input_path, output_path)

    doc = Document(output_path)
    for ti, tbl in enumerate(doc.tables):
        for ri, row in enumerate(tbl.rows):
            cells = row.cells
            if len(cells) < 2:
                continue
            # Skip the header row (R0) of every table — that's where the section
            # title or creative number lives.
            if ri == 0:
                continue
            # Look at the value cell (column 1) — but skip if it contains an SDT
            # (dropdown), since those have their own guidance text inside the
            # SDT placeholder which Word handles separately.
            value_cell = cells[1]
            sdt = value_cell._tc.find('.//' + qn('w:sdt'))
            if sdt is not None:
                continue

            text = cell_text(value_cell)
            if BRACKET_RE.match(text):
                clear_cell(value_cell)

    doc.save(output_path)


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)
    input_path = sys.argv[1]
    if not os.path.exists(input_path):
        print(f'ERROR: file not found: {input_path}')
        sys.exit(1)
    output_path = sys.argv[2] if len(sys.argv) >= 3 else input_path
    strip_placeholders(input_path, output_path)
    print(f'Stripped placeholders: {input_path} -> {output_path}')


if __name__ == '__main__':
    main()
